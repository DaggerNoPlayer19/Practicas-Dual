from datetime import date
from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SKILL_SCRIPTS = Path(
    r"C:\Users\alesh\.codex\plugins\cache\openai-primary-runtime\documents\26.623.12021\skills\documents\scripts"
)
if str(SKILL_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SKILL_SCRIPTS))

from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


ROOT = Path(r"C:\Users\alesh\practica1-autenticacion\practica5")
EVIDENCE_DIR = ROOT / "evidencias"
CAPTURES_DIR = EVIDENCE_DIR / "capturas"
OUTPUT = EVIDENCE_DIR / "Reporte_Practica_05_Sanctum.docx"


def set_run_font(run, *, bold=None, italic=None, size=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def clear_direct_borders(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn("w:pBdr"):
            p_pr.remove(child)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.style = doc.styles["Heading 1"]
    elif level == 2:
        p.style = doc.styles["Heading 2"]
    else:
        p.style = doc.styles["Heading 3"]
    p.add_run(text)
    return p


def add_body(doc, text, *, bold_prefix=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=6, line=1.1)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, after=4, line=1.1)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_caption(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=4, after=2, line=1.0)
    r1 = p.add_run(label)
    set_run_font(r1, bold=True, size=10, color="1F1F1F")
    r2 = p.add_run(text)
    set_run_font(r2, size=10, color="1F1F1F")
    return p


def insert_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=110, after=4, line=1.0)
    r = p.add_run("Pr\u00e1ctica 5")
    set_run_font(r, size=26, bold=False, color="000000")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=6, line=1.0)
    r = p.add_run("Seguridad de APIs con Laravel Sanctum")
    set_run_font(r, size=16, bold=True, color="1F3A5F")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=18, line=1.0)
    r = p.add_run("Reporte de implementaci\u00f3n y evidencias")
    set_run_font(r, size=11, italic=True, color="555555")

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    meta.cell(0, 0).text = "Alumno"
    meta.cell(0, 1).text = "Alejandro Avalos Espinosa"
    meta.cell(1, 0).text = "Materia"
    meta.cell(1, 1).text = "Desarrollo web / Seguridad de aplicaciones"
    meta.cell(2, 0).text = "Fecha"
    meta.cell(2, 1).text = date.today().strftime("%d/%m/%Y")
    meta.cell(3, 0).text = "Proyecto"
    meta.cell(3, 1).text = "practica5"
    widths = column_widths_from_weights([1.875, 4.625])
    apply_table_geometry(meta, widths)
    for row in meta.rows:
        for idx, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                if idx == 0:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_spacing(para, before=0, after=0, line=1.0)
                for run in para.runs:
                    set_run_font(run, bold=(idx == 0), size=11)
            cell.vertical_alignment = 1

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=18, after=0, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Evidencias incluidas al final del documento.")
    set_run_font(r, size=10, color="555555", italic=True)


def add_validation_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Caso"
    hdr[1].text = "Resultado esperado"
    hdr[2].text = "Resultado observado"
    for cell in hdr:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=0, after=0, line=1.0)
            for run in p.runs:
                set_run_font(run, bold=True, size=10)
    rows = [
        (
            "Login con token de solo lectura",
            "El token se emite con la ability ver",
            "El panel muestra abilities: ver",
        ),
        (
            "Intento de alta con token de solo lectura",
            "La acci\u00f3n se bloquea con 403",
            "El sistema responde: tu token no tiene el permiso crear",
        ),
        (
            "Login con token de escritura",
            "El token incluye ver, crear, editar y eliminar",
            "El panel muestra abilities: ver, crear, editar, eliminar",
        ),
        (
            "Alta de producto con token de escritura",
            "El producto se guarda correctamente",
            "La fila aparece en la lista de productos sin recargar",
        ),
    ]
    for item in rows:
        row = table.add_row().cells
        for i, text in enumerate(item):
            row[i].text = text
            for p in row[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_spacing(p, before=0, after=0, line=1.0)
                for run in p.runs:
                    set_run_font(run, size=10)
    apply_table_geometry(table, column_widths_from_weights([2.0, 2.3, 2.2]))


def add_evidence(doc, idx, filename, title, description, max_width=6.0):
    image_path = CAPTURES_DIR / filename
    add_caption(doc, f"Figura {idx}. ", title)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=4, line=1.0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(max_width))
    add_body(doc, description)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        fmt = style.paragraph_format
        fmt.space_before = Pt(before)
        fmt.space_after = Pt(after)
        fmt.line_spacing = 1.1

    add_cover(doc)
    insert_page_break(doc)

    add_heading(doc, "1. Objetivo", 1)
    add_body(
        doc,
        "Implementar autenticaci\u00f3n de APIs con Laravel Sanctum, emitir tokens "
        "con abilities diferenciadas y validar permisos antes de ejecutar operaciones "
        "de escritura sobre productos.",
    )

    add_heading(doc, "2. Desarrollo", 1)
    add_heading(doc, "2.1 Backend", 2)
    add_body(
        doc,
        "Se integr\u00f3 HasApiTokens en el modelo User, se configuraron los endpoints "
        "de login, registro, perfil y logout, y se ajust\u00f3 ProductoController para "
        "permitir solo operaciones compatibles con el permiso del token actual.",
    )
    add_body(
        doc,
        "El backend emite dos conjuntos de abilities: ver para lectura y ver, crear, "
        "editar, eliminar para escritura. Con ello se controla el alcance de cada token "
        "y se responde con mensajes claros cuando falta un permiso.",
    )

    add_heading(doc, "2.2 Frontend", 2)
    add_body(
        doc,
        "La pantalla de acceso permite elegir entre token de solo lectura y token de "
        "escritura. El panel admin muestra las abilities activas para facilitar la "
        "evidencia visual de la sesi\u00f3n.",
    )
    add_body(
        doc,
        "El flujo de creaci\u00f3n, edici\u00f3n y eliminaci\u00f3n se mantiene sin "
        "recargar la p\u00e1gina, de modo que la retroalimentaci\u00f3n del usuario es "
        "inmediata y consistente.",
    )

    add_heading(doc, "2.3 Validaci\u00f3n", 2)
    add_body(
        doc,
        "Se a\u00f1adieron pruebas autom\u00e1ticas para confirmar que los tokens se "
        "generan con las abilities correctas y que un token de solo lectura no puede "
        "crear productos.",
    )
    add_validation_table(doc)

    add_heading(doc, "3. Evidencias", 1)
    add_body(
        doc,
        "Las capturas siguientes documentan el comportamiento de la aplicaci\u00f3n con "
        "distintos tipos de token y muestran el resultado real de la validaci\u00f3n de "
        "permisos.",
    )

    add_evidence(
        doc,
        1,
        "p5_01_login_token_type.png",
        "Formulario de acceso con selector de tipo de token",
        "Se observa la opci\u00f3n para emitir un token de solo lectura o de escritura.",
    )
    add_evidence(
        doc,
        2,
        "p5_02_admin_abilities_read.png",
        "Panel administrativo con token de solo lectura",
        "El panel indica que la sesi\u00f3n activa solo dispone de la ability ver.",
    )
    add_evidence(
        doc,
        3,
        "p5_03_read_token_blocked_create.png",
        "Intento de alta bloqueado por falta de permiso",
        "Al intentar guardar un producto con token de solo lectura aparece el mensaje de "
        "que el token no tiene el permiso crear.",
    )
    add_evidence(
        doc,
        4,
        "p5_04_admin_abilities_write.png",
        "Panel administrativo con token de escritura",
        "La sesi\u00f3n de escritura muestra las abilities ver, crear, editar y eliminar.",
    )
    add_evidence(
        doc,
        5,
        "p5_05_write_token_product_created.png",
        "Alta exitosa de producto con token de escritura",
        "El nuevo producto queda visible en la lista sin necesidad de recargar la vista.",
    )

    add_heading(doc, "4. Resultado final", 1)
    add_body(
        doc,
        "La pr\u00e1ctica qued\u00f3 funcional con autenticaci\u00f3n mediante Sanctum, "
        "control por abilities y evidencia documental completa. El resultado final "
        "cumple con el flujo solicitado: acceso, control de permisos, operaci\u00f3n "
        "correcta de productos y registro de pruebas visuales.",
    )
    add_body(
        doc,
        "Como cierre, el sistema distingue de forma clara entre lectura y escritura, "
        "lo que permite usar el mismo frontend con tokens de diferentes privilegios sin "
        "comprometer la seguridad de las operaciones sensibles.",
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
